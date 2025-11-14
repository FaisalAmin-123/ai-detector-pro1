"""
Multi-Format Content Processor
Handles PDFs, Images, Videos, and Documents

Author: Faisal + Assistant
Date: November 8, 2024
"""

import os
import logging
from typing import Dict, Optional

# PDF processing
try:
    import PyPDF2
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    logging.warning("PDF libraries not installed. Run: pip install PyPDF2 pdfplumber")

# Image OCR
try:
    from PIL import Image
    import pytesseract
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False
    logging.warning("OCR libraries not installed. Run: pip install pillow pytesseract")

# Document processing
try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("DOCX library not installed. Run: pip install python-docx")

# Video processing
try:
    import speech_recognition as sr
    from moviepy.editor import VideoFileClip
    VIDEO_AVAILABLE = True
except ImportError:
    VIDEO_AVAILABLE = False
    logging.warning("Video libraries not installed. Run: pip install SpeechRecognition moviepy")

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class MultiFormatProcessor:
    """
    Extract text from multiple file formats
    """
    
    def __init__(self):
        """Initialize processor"""
        logger.info("Initializing Multi-Format Processor...")
        self.supported_formats = self._check_supported_formats()
        logger.info(f"Supported formats: {', '.join(self.supported_formats)}")
    
    def _check_supported_formats(self):
        """Check which formats are supported based on installed libraries"""
        formats = ['txt']  # Text always supported
        
        if PDF_AVAILABLE:
            formats.extend(['pdf'])
        if OCR_AVAILABLE:
            formats.extend(['jpg', 'jpeg', 'png', 'tiff', 'bmp'])
        if DOCX_AVAILABLE:
            formats.extend(['docx'])
        if VIDEO_AVAILABLE:
            formats.extend(['mp4', 'avi', 'mov', 'mkv'])
        
        return formats
    
    def process_file(self, file_path: str) -> Dict:
        """
        Process any supported file and extract text
        
        Args:
            file_path: Path to file
            
        Returns:
            dict with 'text', 'metadata', 'error'
        """
        if not os.path.exists(file_path):
            return {'error': 'File not found', 'text': None, 'metadata': {}}
        
        # Get file extension
        file_ext = os.path.splitext(file_path)[1].lower().lstrip('.')
        
        if file_ext not in self.supported_formats:
            return {
                'error': f'Unsupported format: {file_ext}',
                'text': None,
                'metadata': {'supported_formats': self.supported_formats}
            }
        
        try:
            # Route to appropriate processor
            if file_ext == 'pdf':
                return self._process_pdf(file_path)
            elif file_ext in ['jpg', 'jpeg', 'png', 'tiff', 'bmp']:
                return self._process_image(file_path)
            elif file_ext == 'docx':
                return self._process_docx(file_path)
            elif file_ext in ['mp4', 'avi', 'mov', 'mkv']:
                return self._process_video(file_path)
            elif file_ext == 'txt':
                return self._process_text(file_path)
            else:
                return {'error': f'No processor for {file_ext}', 'text': None, 'metadata': {}}
        
        except Exception as e:
            logger.error(f"Error processing {file_path}: {e}")
            return {'error': str(e), 'text': None, 'metadata': {}}
    
    def _process_text(self, file_path: str) -> Dict:
        """Process plain text file"""
        logger.info(f"Processing text file: {file_path}")
        
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            text = f.read()
        
        return {
            'text': text,
            'metadata': {
                'format': 'text',
                'length': len(text),
                'word_count': len(text.split())
            },
            'error': None
        }
    
    def _process_pdf(self, file_path: str) -> Dict:
        """Process PDF file"""
        logger.info(f"Processing PDF: {file_path}")
        
        text = ""
        page_count = 0
        
        try:
            # Method 1: Try pdfplumber (better for complex PDFs)
            with pdfplumber.open(file_path) as pdf:
                page_count = len(pdf.pages)
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        except Exception as e:
            logger.warning(f"pdfplumber failed: {e}, trying PyPDF2...")
            
            # Method 2: Fallback to PyPDF2
            try:
                with open(file_path, 'rb') as f:
                    pdf_reader = PyPDF2.PdfReader(f)
                    page_count = len(pdf_reader.pages)
                    for page in pdf_reader.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
            except Exception as e2:
                return {'error': f'PDF extraction failed: {e2}', 'text': None, 'metadata': {}}
        
        if not text.strip():
            return {
                'error': 'PDF appears to be empty or scanned (needs OCR)',
                'text': None,
                'metadata': {'pages': page_count, 'suggestion': 'Try OCR for scanned PDFs'}
            }
        
        return {
            'text': text.strip(),
            'metadata': {
                'format': 'pdf',
                'pages': page_count,
                'length': len(text),
                'word_count': len(text.split())
            },
            'error': None
        }
    
    def _process_image(self, file_path: str) -> Dict:
        """Process image using OCR"""
        logger.info(f"Processing image with OCR: {file_path}")
        
        try:
            image = Image.open(file_path)
            
            # Perform OCR
            text = pytesseract.image_to_string(image)
            
            if not text.strip():
                return {
                    'error': 'No text detected in image',
                    'text': None,
                    'metadata': {'format': 'image', 'size': image.size}
                }
            
            return {
                'text': text.strip(),
                'metadata': {
                    'format': 'image',
                    'dimensions': image.size,
                    'mode': image.mode,
                    'length': len(text),
                    'word_count': len(text.split())
                },
                'error': None
            }
        
        except Exception as e:
            return {'error': f'OCR failed: {e}', 'text': None, 'metadata': {}}
    
    def _process_docx(self, file_path: str) -> Dict:
        """Process Word document"""
        logger.info(f"Processing DOCX: {file_path}")
        
        try:
            doc = docx.Document(file_path)
            
            # Extract text from paragraphs
            text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
            
            if not text.strip():
                return {'error': 'DOCX appears to be empty', 'text': None, 'metadata': {}}
            
            return {
                'text': text.strip(),
                'metadata': {
                    'format': 'docx',
                    'paragraphs': len(doc.paragraphs),
                    'length': len(text),
                    'word_count': len(text.split())
                },
                'error': None
            }
        
        except Exception as e:
            return {'error': f'DOCX processing failed: {e}', 'text': None, 'metadata': {}}
    
    def _process_video(self, file_path: str) -> Dict:
        """Process video by extracting audio and transcribing"""
        logger.info(f"Processing video: {file_path}")
        
        try:
            # Extract audio
            video = VideoFileClip(file_path)
            audio_path = file_path.replace(os.path.splitext(file_path)[1], '_temp.wav')
            video.audio.write_audiofile(audio_path, logger=None)
            
            # Transcribe audio
            recognizer = sr.Recognizer()
            with sr.AudioFile(audio_path) as source:
                audio_data = recognizer.record(source)
                text = recognizer.recognize_google(audio_data)
            
            # Cleanup temp audio
            if os.path.exists(audio_path):
                os.remove(audio_path)
            
            return {
                'text': text,
                'metadata': {
                    'format': 'video',
                    'duration': video.duration,
                    'fps': video.fps,
                    'length': len(text),
                    'word_count': len(text.split())
                },
                'error': None
            }
        
        except Exception as e:
            return {'error': f'Video processing failed: {e}', 'text': None, 'metadata': {}}


# Test
if __name__ == "__main__":
    print("=== Multi-Format Processor Test ===\n")
    
    processor = MultiFormatProcessor()
    
    print(f"Supported formats: {processor.supported_formats}\n")
    
    # Test with a sample text file
    test_text = "This is a test file for AI detection."
    test_file = "test_sample.txt"
    
    with open(test_file, 'w') as f:
        f.write(test_text)
    
    print("Testing text file processing...")
    result = processor.process_file(test_file)
    
    if result['error']:
        print(f"❌ Error: {result['error']}")
    else:
        print(f"✅ Success!")
        print(f"Text: {result['text']}")
        print(f"Metadata: {result['metadata']}")
    
    # Cleanup
    os.remove(test_file)
    
    print("\n=== Test Complete ===")